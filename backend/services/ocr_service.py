# /services/ocr_service.py

import os
import re
import json
import uuid
from pathlib import Path
from typing import List, Tuple

import numpy as np
import cv2
import torch
from PIL import ImageFilter, ImageEnhance, Image
from pdf2image import convert_from_path
from scipy import ndimage
from scipy.signal import find_peaks
from tqdm import tqdm

from dataclasses import dataclass, field
import re
import docx

from transformers import TrOCRProcessor, VisionEncoderDecoderModel

from schemas.document import RecognitionResult, TranscriptData
from core.config import settings

_token_re = re.compile(r"[А-Яа-яA-Za-zЁёІіѢѣѲѳѴѵ]+", flags=re.UNICODE)

# --- Опциональные зависимости ---
try:
    from mmocr.apis import TextDetInferencer as MMOCRTextDetInferencer
    _HAS_MMOCR = True
except ImportError:
    MMOCRTextDetInferencer = None
    _HAS_MMOCR = False

try:
    import kenlm
    _HAS_KENLM = True
except ImportError:
    kenlm = None
    _HAS_KENLM = False

try:
    import pymorphy2
    _HAS_PYMORPHY = True
except ImportError:
    pymorphy2 = None
    _HAS_PYMORPHY = False

# --- Конфигурация и глобальные объекты ---

# Хранилище для прогресса (в реальном проекте - Celery/Redis)
progress_status = {}

# Глобальные настройки из ocr.py
DEVICE = torch.device("cuda" if torch.cuda.is_available() else "cpu")
USE_DBNET = True
DBNET_MODEL_NAME = os.getenv("DBNET_MODEL_NAME", "dbnet_resnet18_fpnc_1200e_icdar2015")
USE_DEWARP = True
DEWARPNET_ENABLED = os.getenv("DEWARPNET_ENABLED", "0") == "1"
DEWARPNET_WC = os.getenv("DEWARPNET_WC_PATH", None)
DEWARPNET_BM = os.getenv("DEWARPNET_BM_PATH", None)
SAUVOLA_ON_LOW_CONTRAST = True
SAUVOLA_WINDOW = 25
SAUVOLA_K = 0.2
LOW_CONTRAST_THRESHOLD = 20.0
BEAM_SIZE = int(os.getenv("BEAM_SIZE", "5"))
N_BEST = int(os.getenv("N_BEST", "5"))
MAX_GEN_LEN = int(os.getenv("MAX_GEN_LEN", "128"))
LAMBDA_MODEL = float(os.getenv("LAMBDA_MODEL", "1.0"))
LAMBDA_LM = float(os.getenv("LAMBDA_LM", "0.15"))
LAMBDA_OOV = float(os.getenv("LAMBDA_OOV", "0.5"))
LENGTH_NORM = float(os.getenv("LENGTH_NORM", "0.0"))
KENLM_MODEL_PATH = os.getenv("KENLM_ARPA", None)
ALPHAVIT_DOCX_PATH = "/backend/Alphabet.docx"
FUND_DOCX_PATH = "/backend/F203.docx"

print(f"Device: {DEVICE}")

# --- Инициализация моделей (ленивая загрузка) ---
_processor = None
_model = None
_morph_analyzer = None
_lm_model = None
_dbnet_inferencer = None
_token_re = re.compile(r"[А-Яа-яA-Za-zЁё]+", flags=re.UNICODE)

def get_processor():
    global _processor
    if _processor is None:
        _processor = TrOCRProcessor.from_pretrained('kazars24/trocr-base-handwritten-ru')
    return _processor

def get_model():
    global _model
    if _model is None:
        _model = VisionEncoderDecoderModel.from_pretrained('kazars24/trocr-base-handwritten-ru').to(DEVICE)
        _model.eval()
    return _model

def get_morph_analyzer():
    global _morph_analyzer
    if _morph_analyzer is None and _HAS_PYMORPHY:
        _morph_analyzer = pymorphy2.MorphAnalyzer(lang='ru')
    return _morph_analyzer

def get_lm_model():
    global _lm_model
    if _lm_model is None and _HAS_KENLM and KENLM_MODEL_PATH and os.path.exists(KENLM_MODEL_PATH):
        _lm_model = kenlm.LanguageModel(KENLM_MODEL_PATH)
    return _lm_model

def get_dbnet_inferencer():
    global _dbnet_inferencer
    if _dbnet_inferencer is None and _HAS_MMOCR:
        try:
            _dbnet_inferencer = MMOCRTextDetInferencer(model=DBNET_MODEL_NAME, device=str(DEVICE))
        except Exception:
            try:
                _dbnet_inferencer = MMOCRTextDetInferencer(model='DBNet', device=str(DEVICE))
            except Exception:
                pass
    return _dbnet_inferencer


# --- Вспомогательные функции из ocr.py ---

def preprocess_for_segmentation(image: Image.Image) -> Image.Image:
    enhancer = ImageEnhance.Contrast(image)
    image = enhancer.enhance(1.2)
    image = image.filter(ImageFilter.GaussianBlur(radius=0.3))
    image = image.filter(ImageFilter.UnsharpMask(radius=1, percent=50, threshold=3))
    return image

def sauvola_binarize_pil(image: Image.Image, window_size: int = SAUVOLA_WINDOW, k: float = SAUVOLA_K) -> Image.Image:
    if image.mode != 'L':
        grayscale = image.convert('L')
    else:
        grayscale = image
    img = np.array(grayscale).astype(np.float32)
    mean = cv2.boxFilter(img, -1, (window_size, window_size))
    sqmean = cv2.boxFilter(img**2, -1, (window_size, window_size))
    var = np.maximum(0.0, sqmean - mean**2)
    std = np.sqrt(var)
    thresh = mean * (1 + k * ((std / 128.0) - 1))
    binary = (img > thresh).astype(np.uint8) * 255
    return Image.fromarray(binary)

def page_contrast_std(image: Image.Image) -> float:
    img = image.convert('L') if image.mode != 'L' else image
    return float(np.array(img).astype(np.float32).std())

def _dewarp_perspective_basic(pil_img: Image.Image) -> Image.Image:
    img = np.array(pil_img.convert('RGB'))
    g = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    g = cv2.GaussianBlur(g, (5,5), 0)
    edges = cv2.Canny(g, 50, 150)
    contours, _ = cv2.findContours(edges, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)
    H, W = g.shape
    best = max([cnt for cnt in contours if len(cv2.approxPolyDP(cnt, 0.02 * cv2.arcLength(cnt, True), True)) == 4 and cv2.isContourConvex(cv2.approxPolyDP(cnt, 0.02 * cv2.arcLength(cnt, True), True))], key=cv2.contourArea, default=None)

    if best is None or cv2.contourArea(best) < 0.2 * H * W:
        return pil_img

    pts = best.reshape(-1, 2).astype(np.float32)
    s = pts.sum(axis=1)
    diff = np.diff(pts, axis=1).ravel()
    ordered_pts = np.array([pts[np.argmin(s)], pts[np.argmin(diff)], pts[np.argmax(s)], pts[np.argmax(diff)]], dtype=np.float32)

    widthA = np.linalg.norm(ordered_pts[2] - ordered_pts[3])
    widthB = np.linalg.norm(ordered_pts[1] - ordered_pts[0])
    heightA = np.linalg.norm(ordered_pts[1] - ordered_pts[2])
    heightB = np.linalg.norm(ordered_pts[0] - ordered_pts[3])
    maxW, maxH = int(max(widthA, widthB)), int(max(heightA, heightB))

    dst = np.array([[0, 0], [maxW - 1, 0], [maxW - 1, maxH - 1], [0, maxH - 1]], dtype=np.float32)
    M = cv2.getPerspectiveTransform(ordered_pts, dst)
    warped = cv2.warpPerspective(img, M, (maxW, maxH), flags=cv2.INTER_CUBIC)
    return Image.fromarray(warped)

def dewarp_image(image: Image.Image) -> Image.Image:
    if not USE_DEWARP: return image
    # Здесь должен быть вызов DewarpNet, пока используется базовый фоллбэк
    return _dewarp_perspective_basic(image)

def _group_boxes_to_lines(boxes: List[np.ndarray], y_tol: int = 10) -> List[Tuple[int,int,int,int]]:
    if not boxes: return []
    entries = sorted([( (p[:,1].min() + p[:,1].max()) / 2, p) for p in boxes])
    lines, current_line, current_y = [], [], None
    for y_center, poly in entries:
        if current_y is None or abs(y_center - current_y) <= y_tol:
            current_line.append(poly)
            current_y = y_center if current_y is None else np.mean([current_y, y_center])
        else:
            if current_line: lines.append(np.concatenate(current_line))
            current_line, current_y = [poly], y_center
    if current_line: lines.append(np.concatenate(current_line))
    return [(int(l[:,0].min()), int(l[:,1].min()), int(l[:,0].max()), int(l[:,1].max())) for l in lines]


def segment_lines(image: Image.Image, min_line_height: int = 10) -> List[Tuple[int,int,int,int]]:
    if USE_DBNET and get_dbnet_inferencer():
        np_rgb = np.array(image.convert('RGB'))
        out = get_dbnet_inferencer()(np_rgb, return_vis=False)
        polys = out['predictions'][0].get('polygons', [])
        polys = [np.array(p).reshape(-1, 2) for p in polys if (np.array(p).reshape(-1, 2)[:,1].max() - np.array(p).reshape(-1, 2)[:,1].min()) >= min_line_height]
        if polys: return _group_boxes_to_lines(polys, y_tol=max(8, min_line_height//2))

    # Fallback to projection
    processed_img = preprocess_for_segmentation(image)
    binary_img = sauvola_binarize_pil(processed_img)
    img_array = 255 - np.array(binary_img) if np.mean(np.array(binary_img)) > 127 else np.array(binary_img)
    proj = ndimage.gaussian_filter1d(np.sum(img_array > 0, axis=1), sigma=1)
    gaps, _ = find_peaks(np.max(proj) - proj, height=np.max(proj) * 0.5, distance=min_line_height)
    boundaries = sorted([0] + list(gaps) + [image.height])
    boxes = []
    for i in range(len(boundaries) - 1):
        y0, y1 = boundaries[i], boundaries[i+1]
        if y1 - y0 < min_line_height or np.max(proj[y0:y1]) < np.max(proj) * 0.1: continue
        text_indices = np.where(proj[y0:y1] > np.max(proj[y0:y1]) * 0.1)[0]
        if len(text_indices) == 0: continue
        boxes.append((0, max(0, y0 + text_indices[0] - 2), image.width, min(image.height, y0 + text_indices[-1] + 2)))
    return boxes

def split_double_page(image: Image.Image) -> Tuple[Image.Image, Image.Image]:
    w, h = image.size
    return image.crop((0, 0, w // 2, h)), image.crop((w // 2, 0, w, h))

def _count_oov_tokens(text: str) -> int:
    if not get_morph_analyzer(): return 0
    return sum(1 for tok in _token_re.findall(text) if not get_morph_analyzer().parse(tok) or get_morph_analyzer().parse(tok)[0].score < 0.01)

def _lm_score(text: str) -> float:
    return float(get_lm_model().score(text, bos=True, eos=True)) if get_lm_model() else 0.0

def predict_text_from_line_image(line_image: Image.Image) -> Tuple[str, float]:
    processor, model = get_processor(), get_model()
    inputs = processor(images=line_image.convert('RGB'), return_tensors="pt").to(DEVICE)
    with torch.no_grad():
        gen_out = model.generate(
            inputs.pixel_values, max_length=MAX_GEN_LEN, num_beams=BEAM_SIZE,
            num_return_sequences=N_BEST, early_stopping=True,
            return_dict_in_generate=True, output_scores=True
        )
    texts = processor.batch_decode(gen_out.sequences, skip_special_tokens=True)
    model_scores = gen_out.sequences_scores.detach().cpu().numpy().tolist() if hasattr(gen_out, "sequences_scores") else [0.0] * len(texts)

    best_text, best_total = "", -1e9
    for i, t in enumerate(texts):
        total = LAMBDA_MODEL * model_scores[i] + LAMBDA_LM * _lm_score(t) - LAMBDA_OOV * _count_oov_tokens(t)
        if total > best_total:
            best_total, best_text = total, t
    
    # Confidence as normalized model score
    confidence = float(np.exp(min(model_scores) / len(best_text))) if best_text else 0.0
    return best_text, confidence

# --- Основная функция обработки, адаптированная под сервис ---

def process_document(file_id: uuid.UUID, file_path: str):
    """
    Реальная функция нормализации и распознавания документа.
    """
    print(f"Начало обработки файла: {file_path}")
    progress_status[file_id] = {"status": "starting", "progress": 0}

    try:
        # --- 1. Загрузка и подготовка изображений ---
        progress_status[file_id] = {"status": "loading", "progress": 10}
        
        fpath = Path(file_path)
        if fpath.suffix.lower() == '.pdf':
            pages = convert_from_path(str(fpath), dpi=300)
        else:
            pages = [Image.open(fpath)]

        all_recognized_words = []
        page_texts = []
        total_lines = 0

        progress_status[file_id] = {"status": "loading", "progress": 10}
        fpath = Path(file_path)

        # NEW: пути к DOCX из окружения
        alphavit_docx = os.getenv("ALPHAVIT_DOCX_PATH", None)  # NEW
        fund_docx = os.getenv("FUND_DOCX_PATH", None)          # NEW
        alpha_rules = load_alphabet_from_docx(alphavit_docx)   # NEW
        glossary = load_fund_glossary(fund_docx, alpha_rules)  # NEW

        if fpath.suffix.lower() == '.pdf':
            pages = convert_from_path(str(fpath), dpi=300)
        else:
            pages = [Image.open(fpath)]
                
        # --- 2. Обработка каждой страницы ---
        for page_idx, page_img in enumerate(pages):
            progress_status[file_id] = {"status": "processing_page", "progress": 20 + 60 * (page_idx / len(pages))}
            
            # Разделение разворота
            left, right = split_double_page(page_img)
            
            for side_img, side_name in [(left, 'left'), (right, 'right')]:
                # Деварпинг
                dewarped_img = dewarp_image(side_img)
                
                # Сегментация строк
                lines_coords = segment_lines(dewarped_img, min_line_height=30)
                if not lines_coords: continue
                total_lines += len(lines_coords)

                # Распознавание каждой строки
                for i, coords in enumerate(lines_coords):
                    progress_status[file_id] = {"status": "recognizing", "progress": 25 + 60 * (page_idx / len(pages)) + 30 * (i / len(lines_coords))}
                    x0, y0, x1, y1 = coords
                    pad = max(2, (y1 - y0) // 20)
                    line_img = dewarped_img.crop((max(0, x0 - 5), max(0, y0 - pad), min(dewarped_img.width, x1 + 5), min(dewarped_img.height, y1 + pad)))
                    
                    line_text, confidence = predict_text_from_line_image(line_img)
                    if not line_text.strip(): 
                        continue
                    norm_text = normalize_and_correct_line(line_text, glossary, alpha_rules)
                    page_texts.append(norm_text)

                    for word in norm_text.split():
                        all_recognized_words.append(
                            TranscriptData(text=word, coordinates=[x0, y0, x1, y1], confidence=round(confidence, 3)))

        # --- 3. Формирование и сохранение результата ---
        progress_status[file_id] = {"status": "saving", "progress": 95}
        
        # WER можно рассчитать, если есть ground truth. Пока ставим 0.
        result = RecognitionResult(
            wer=0.0,
            recognized_words=all_recognized_words,
            extracted_attributes={} # Заглушка для NER
        )

        transcript_filename = f"{file_id}_transcript.json"
        transcript_path = os.path.join(settings.TRANSCRIPTS_DIRECTORY, transcript_filename)
        transcript_path = transcript_path.replace('\\', '/')
        with open(transcript_path, "w", encoding="utf-8") as f:
            f.write(result.model_dump_json(indent=4))

        print(f"Обработка файла {file_path} завершена. Найдено {total_lines} строк. Результат в {transcript_path}")
        
    except Exception as e:
        print(f"Ошибка при обработке {file_path}: {e}")
        progress_status[file_id] = {"status": "error", "progress": 100, "message": str(e)}
        # В реальном приложении здесь должен быть более сложный механизм обработки ошибок
        return None, -1
    finally:
        if file_id in progress_status and progress_status[file_id].get("status") != "error":
            progress_status[file_id] = {"status": "completed", "progress": 100}

    return transcript_path, result.wer


# --- Mock-функция остается для обратной совместимости или тестов ---

def process_document_mock(file_id: uuid.UUID, file_path: str):
    """
    Имитация процесса для тестирования API.
    """
    # ... (код mock-функции остается без изменений) ...
    pass

def get_progress(file_id: uuid.UUID):
    return progress_status.get(file_id, {"status": "not_found", "progress": 0})


@dataclass
class AlphabetRules:
    char_map: dict[str, str] = field(default_factory=dict)
    remove_final_er: bool = False
    replace_ago_ego: bool = False
    extra_word_rules: list[tuple[re.Pattern, str]] = field(default_factory=list)

# финальный ъ
_RE_FINAL_ER = re.compile(r"(\b\w*?[бвгджзклмнпрстфхцчшщ])ъ\b", flags=re.IGNORECASE | re.UNICODE)

def _parse_pairs_from_text(txt: str) -> dict[str, str]:
    pairs = {}
    # Поддержка форматов: "ѣ -> е", "Ѣ→Е", "і = и", "ѳ=ф", "ѵ->и"
    for a, b in re.findall(r"([А-Яа-яЁёІіѢѣѲѳѴѵ])\s*[-=→>]+\s*([А-Яа-яЁёІіѢѣѲѳѴѵ])", txt):
        pairs[a] = b
        # автодобавление верхнего регистра
        if a.upper() != a and b.upper() != b:
            pairs[a.upper()] = b.upper()
    return pairs

def load_alphabet_from_docx(alphavit_docx_path: str | None) -> AlphabetRules:
    rules = AlphabetRules()
    if not alphavit_docx_path or not os.path.exists(alphavit_docx_path) or not docx:
        return rules
    d = docx.Document(alphavit_docx_path)
    # 1) пары из таблиц
    for tbl in d.tables:
        for row in tbl.rows:
            if len(row.cells) >= 2:
                a = row.cells[0].text.strip()
                b = row.cells[1].text.strip()
                if a and b:
                    rules.char_map[a[:1]] = b[:1]
                    if a[:1].upper() != a[:1] and b[:1].upper() != b[:1]:
                        rules.char_map[a[:1].upper()] = b[:1].upper()
    # 2) пары из текста и эвристики по правилам
    full = "\n".join(p.text for p in d.paragraphs)
    rules.char_map.update(_parse_pairs_from_text(full))
    low = full.lower()
    # включить удаление финального ъ, если упоминается
    if "твердый знак" in low or "ъ" in low:
        rules.remove_final_er = True
    # включить -аго/-яго → -ого/-его, если упомянуто
    if "-аго" in low or "-яго" in low:
        rules.replace_ago_ego = True
    if re.search(r"окончани[ея].*?-аго.*?-ого", low):
        rules.extra_word_rules.append((re.compile(r"(\w+?)аго\b"), r"\1ого"))
    if re.search(r"окончани[ея].*?-яго.*?-его", low):
        rules.extra_word_rules.append((re.compile(r"(\w+?)яго\b"), r"\1его"))
    return rules

def normalize_pre_reform(s: str, rules: AlphabetRules) -> str:
    t = "".join(rules.char_map.get(ch, ch) for ch in s)
    if rules.remove_final_er:
        t = _RE_FINAL_ER.sub(lambda m: m.group(1), t)
    if rules.replace_ago_ego:
        t = re.sub(r"(\w+?)аго\b", r"\1ого", t)
        t = re.sub(r"(\w+?)яго\b", r"\1его", t)
    for pat, repl in rules.extra_word_rules:
        t = pat.sub(repl, t)
    return t

def _build_glossary_from_text(txt: str, rules: AlphabetRules) -> set[str]:
    tokens = re.findall(r"[А-Яа-яЁёІіѢѣѲѳѴѵ-]+", txt)
    norm = { normalize_pre_reform(w, rules).lower() for w in tokens if len(w) >= 3 }
    return {w for w in norm if len(w) >= 3}

def load_fund_glossary(docx_path: str | None, rules: AlphabetRules) -> set[str]:
    if not docx_path or not os.path.exists(docx_path) or not docx:
        return set()
    d = docx.Document(docx_path)
    txt = "\n".join(p.text for p in d.paragraphs)
    return _build_glossary_from_text(txt, rules)

def _weighted_edit_distance(a: str, b: str, rules: AlphabetRules) -> int:
    # Сниженные штрафы для пар из твоего «Алфавита»
    pairs = set()
    for k, v in rules.char_map.items():
        pairs.add((v.lower(), v.lower()))
        pairs.add((k.lower(), v.lower()))
        pairs.add((v.lower(), k.lower()))
    pairs.update({('е','ё'),('ё','е')})
    def sub_cost(ca, cb):
        ca, cb = ca.lower(), cb.lower()
        return 0 if (ca == cb) or ((ca,cb) in pairs) else 1
    la, lb = len(a), len(b)
    dp = [[0]*(lb+1) for _ in range(la+1)]
    for i in range(la+1): dp[i][0] = i
    for j in range(lb+1): dp[0][j] = j
    for i in range(1, la+1):
        for j in range(1, lb+1):
            dp[i][j] = min(
                dp[i-1][j] + (0 if (a[i-1]=='ъ' and rules.remove_final_er) else 1),
                dp[i][j-1] + 1,
                dp[i-1][j-1] + sub_cost(a[i-1], b[j-1])
            )
    return dp[la][lb]

def _autocorrect_token(tok: str, glossary: set[str], rules: AlphabetRules, max_dist: int = 2) -> str:
    if not tok:
        return tok
    ntok = normalize_pre_reform(tok, rules)
    low = ntok.lower()
    if low in glossary:
        return ntok
    best, bestd = ntok, max_dist+1
    for w in glossary:
        d = _weighted_edit_distance(low, w, rules)
        if d < bestd:
            best, bestd = w, d
            if d == 0: break
    if bestd <= max_dist:
        return best.capitalize() if tok[:1].isupper() else best
    return ntok

def normalize_and_correct_line(text: str, glossary: set[str], rules: AlphabetRules) -> str:
    parts = re.findall(r"[А-Яа-яЁёІіѢѣѲѳѴѵ]+|[^А-Яа-яЁёІіѢѣѲѳѴѵ]+", text)
    out = []
    for p in parts:
        if re.match(r"[А-Яа-яЁёІіѢѣѲѳѴѵ]+", p):
            out.append(_autocorrect_token(p, glossary, rules))
        else:
            out.append(p)
    return "".join(out)
